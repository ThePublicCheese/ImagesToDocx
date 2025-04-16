from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

#example usage
imageFolder = './'  #path to your image folder (rn its just current directory)
outputPath = './output.docx'  #path to save the Word document

#default values for a 184 image cheatsheet is 7,11,1.15 and the entire sheet will be filled

imagesPerRow=     7 #change for a different amount of images per row
imagesPerColumn=  11 #change for a different amount of images per column
imageWidthInch=   1.15 #change for a different image width (overall size)


def remove_table_cell_margins(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')

    for margin_type in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{margin_type}')
        margin.set(qn('w:w'), "0")
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)

    tblPr.append(tblCellMar)

def add_images_to_docx(imageFolder, outputPath, imagesPerRow, imagesPerColumn, imageWidthInch):
    doc = Document()

    # Set page margins to 0.17 inches
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.17)
        section.bottom_margin = Inches(0.17)
        section.left_margin = Inches(0.17)
        section.right_margin = Inches(0.17)

    images = [f for f in os.listdir(imageFolder) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))]
    total_images = len(images)

    # Calculate total number of rows needed
    num_rows = (total_images + imagesPerRow - 1) // imagesPerRow  # ceiling division

    # Create one big table
    table = doc.add_table(rows=num_rows, cols=imagesPerRow)
    remove_table_cell_margins(table)

    for img_idx, image_name in enumerate(images):
        row_idx = img_idx // imagesPerRow
        col_idx = img_idx % imagesPerRow

        image_path = os.path.join(imageFolder, image_name)
        cell = table.rows[row_idx].cells[col_idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(imageWidthInch))

    doc.save(outputPath)

add_images_to_docx(imageFolder, outputPath, imagesPerRow, imagesPerColumn, imageWidthInch)
print("*DING* Your sheet is ready sire 🙇")
